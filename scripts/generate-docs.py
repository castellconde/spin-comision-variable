#!/usr/bin/env python3
"""
generate-docs.py — Convierte el Markdown que produce el agente doc-writer
en un .docx con header FEMSA Comercio, footer con paginación, y estilos
corporativos similares a los .dotx oficiales (FCTI_CNF_*).

Uso:
    python scripts/generate-docs.py docs/Esp_Tecnica_v0.3.md
    python scripts/generate-docs.py docs/Runbook_v0.2.md --doc-type runbook
    python scripts/generate-docs.py docs/Esp_Tecnica_v0.3.md --output out.docx

El tipo de documento se infiere del nombre del archivo:
    Esp_Tecnica_*       → "Especificación Técnica" / FCTI_CNF_Especificación Técnica
    Runbook_*           → "Runbook" / FCTI_CNF_Runbook
    Matriz_Pruebas_*    → "Matriz de Pruebas Unitarias"
    ARA_*               → "Análisis de Riesgo Aplicativo"
    MDA_*               → "Modelado de Amenazas"
    Principios_*        → "Principios de Seguridad"

Dependencias:
    pip install python-docx

El logo de FEMSA Comercio vive en .claude/standards/femsa_comercio_logo.png
(extraído del template oficial FCTI_CNF_Especificación_Técnica_Rev_5.dotx).
"""

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: falta python-docx. Instala con: pip install python-docx")


# ─────────────────────────────────────────────────────────────────────────────
# Configuración de tipo de documento
# ─────────────────────────────────────────────────────────────────────────────
DOC_TYPE_CONFIG = {
    "esp_tecnica": {
        "filename_pattern": r"Esp_Tecnica",
        "header_title": "Especificación Técnica",
        "header_code": "FCTI_CNF_Especificación Técnica",
        "header_revision": "REVISIÓN 5",
        "header_revision_date": "29 / 07 / 2016",
        "header_elaboration": "04 / 04 / 2012",
    },
    "runbook": {
        "filename_pattern": r"Runbook",
        "header_title": "Runbook",
        "header_code": "FCTI_CNF_Runbook",
        "header_revision": "REVISIÓN 2",
        "header_revision_date": "",
        "header_elaboration": "20 / 10 / 2014",
    },
    "matriz_pruebas": {
        "filename_pattern": r"Matriz_Pruebas",
        "header_title": "Matriz de Pruebas Unitarias",
        "header_code": "FCTI_CNF_MatrizPruebas",
        "header_revision": "REVISIÓN 1",
        "header_revision_date": "",
        "header_elaboration": "",
    },
    "ara": {
        "filename_pattern": r"ARA",
        "header_title": "Análisis de Riesgo Aplicativo",
        "header_code": "FCTI_CNF_ARA",
        "header_revision": "REVISIÓN 1",
        "header_revision_date": "",
        "header_elaboration": "",
    },
    "mda": {
        "filename_pattern": r"MDA",
        "header_title": "Modelado de Amenazas",
        "header_code": "FCTI_CNF_MDA",
        "header_revision": "REVISIÓN 1",
        "header_revision_date": "",
        "header_elaboration": "",
    },
    "principios_seguridad": {
        "filename_pattern": r"Principios_Seguridad",
        "header_title": "Principios de Seguridad",
        "header_code": "FCTI_CNF_PrincipiosSeguridad",
        "header_revision": "REVISIÓN 1",
        "header_revision_date": "",
        "header_elaboration": "",
    },
}


def infer_doc_type(md_filename: str) -> str:
    """Determina el tipo de documento a partir del nombre del archivo .md"""
    name = Path(md_filename).name
    for doc_type, config in DOC_TYPE_CONFIG.items():
        if re.search(config["filename_pattern"], name, re.IGNORECASE):
            return doc_type
    return "esp_tecnica"  # default


# ─────────────────────────────────────────────────────────────────────────────
# Configuración de estilos corporativos
# ─────────────────────────────────────────────────────────────────────────────
FONT_NAME = "Calibri"  # Fuente cuerpo en los .docx oficiales (verificado en EspTécnica_Planogramas)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY_HEADER = RGBColor(0x40, 0x40, 0x40)
COLOR_TABLE_BORDER = RGBColor(0x80, 0x80, 0x80)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers para manipular XML de python-docx
# ─────────────────────────────────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    """Aplica bordes a una celda. kwargs: top, bottom, left, right (each a dict with sz, val, color)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            for attr, val in kwargs[edge].items():
                border.set(qn(f"w:{attr}"), str(val))
            existing = tcBorders.find(qn(f"w:{edge}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(border)


def set_cell_background(cell, color_hex):
    """Aplica fondo de color a una celda. color_hex sin #."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_page_number(paragraph):
    """Agrega campo de número de página + total al párrafo (para footer)."""
    run1 = paragraph.add_run("Página ")
    run1.font.size = Pt(9)

    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run2 = paragraph.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run2.font.size = Pt(9)

    paragraph.add_run(" de ").font.size = Pt(9)

    # NUMPAGES field
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    run3 = paragraph.add_run()
    run3._r.append(fldChar3)
    run3._r.append(instrText2)
    run3._r.append(fldChar4)
    run3.font.size = Pt(9)


# ─────────────────────────────────────────────────────────────────────────────
# Construcción del header oficial FEMSA Comercio
# ─────────────────────────────────────────────────────────────────────────────
def build_header(section, doc_type: str, logo_path: Path):
    """
    Construye el header oficial replicando la estructura de las plantillas FCTI:
    una tabla con 3 columnas: [logo FEMSA Comercio] | [Sistema de Trabajo TI / <tipo>] | [CÓDIGO / REVISIÓN / ELABORACIÓN].
    """
    config = DOC_TYPE_CONFIG[doc_type]
    header = section.header

    # Limpiar paragraphs default
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # Crear tabla del header (3 columnas, 3 filas para replicar el patrón Planogramas)
    table = header.add_table(rows=3, cols=3, width=Cm(17))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Anchos de columna
    col_widths = [Cm(3.5), Cm(10.0), Cm(3.5)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    # Borders ligeros en todas las celdas
    border_def = {"sz": "4", "val": "single", "color": "808080"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_def, bottom=border_def, left=border_def, right=border_def)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Logo en celda (0,0) — merge verticalmente las 3 filas
    cell_logo = table.cell(0, 0)
    cell_logo.merge(table.cell(2, 0))
    # Limpiar paragraph default
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.text = ""
    if logo_path.exists():
        run = p_logo.add_run()
        run.add_picture(str(logo_path), width=Cm(2.8))
    else:
        # Fallback si no hay logo
        run = p_logo.add_run("FEMSA\nCOMERCIO")
        run.font.bold = True
        run.font.size = Pt(10)

    # Centro: "Sistema de Trabajo TI" en (0,1), título tipo doc en (1,1), vacío en (2,1)
    cell_center_top = table.cell(0, 1)
    cell_center_top.merge(table.cell(2, 1))
    p_center = cell_center_top.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.text = ""

    run1 = p_center.add_run("Sistema de Trabajo TI\n")
    run1.font.bold = True
    run1.font.size = Pt(11)
    run1.font.name = FONT_NAME

    run2 = p_center.add_run(config["header_title"])
    run2.font.bold = True
    run2.font.size = Pt(13)
    run2.font.name = FONT_NAME

    # Derecha: CÓDIGO + REVISIÓN + ELABORACIÓN en 3 filas
    cell_code = table.cell(0, 2)
    p = cell_code.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"CÓDIGO:\n{config['header_code']}")
    run.font.size = Pt(8)
    run.font.name = FONT_NAME

    cell_rev = table.cell(1, 2)
    p = cell_rev.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rev_text = config["header_revision"]
    if config["header_revision_date"]:
        rev_text += f"\n{config['header_revision_date']}"
    run = p.add_run(rev_text)
    run.font.size = Pt(8)
    run.font.name = FONT_NAME

    cell_elab = table.cell(2, 2)
    p = cell_elab.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elab_text = f"ELABORACIÓN: {config['header_elaboration']}" if config['header_elaboration'] else ""
    run = p.add_run(elab_text)
    run.font.size = Pt(8)
    run.font.name = FONT_NAME


def build_footer(section):
    """Footer con 'Página X de Y' centrado."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p)


# ─────────────────────────────────────────────────────────────────────────────
# Parser de Markdown simple — line-by-line
# ─────────────────────────────────────────────────────────────────────────────
def parse_markdown_to_blocks(md_text: str):
    """
    Parsea Markdown a una lista de bloques: dict con type + content.
    Soporta: heading (h1-h6), paragraph, table, list, code, hr.
    """
    blocks = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Heading
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue

        # Code block
        if line.startswith("```"):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines)})
            continue

        # Table (must start with | and next line has separator)
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1]):
            table_rows = []
            # parse header
            header_row = [c.strip() for c in line.strip("|").split("|")]
            table_rows.append(header_row)
            i += 2  # skip separator
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip("|").split("|")]
                table_rows.append(row)
                i += 1
            blocks.append({"type": "table", "rows": table_rows})
            continue

        # Horizontal rule
        if re.match(r"^-{3,}\s*$", line) or re.match(r"^\*{3,}\s*$", line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # List item (- or *)
        m = re.match(r"^(\s*)([-*])\s+(.+)$", line)
        if m:
            indent_level = len(m.group(1)) // 2
            list_items = [{"text": m.group(3), "indent": indent_level}]
            i += 1
            while i < len(lines):
                m2 = re.match(r"^(\s*)([-*])\s+(.+)$", lines[i])
                if m2:
                    list_items.append({"text": m2.group(3), "indent": len(m2.group(1)) // 2})
                    i += 1
                elif lines[i].strip() == "":
                    break
                else:
                    break
            blocks.append({"type": "list", "ordered": False, "items": list_items})
            continue

        # Numbered list
        m = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        if m:
            list_items = [{"text": m.group(3), "indent": len(m.group(1)) // 2}]
            i += 1
            while i < len(lines):
                m2 = re.match(r"^(\s*)(\d+)\.\s+(.+)$", lines[i])
                if m2:
                    list_items.append({"text": m2.group(3), "indent": len(m2.group(1)) // 2})
                    i += 1
                elif lines[i].strip() == "":
                    break
                else:
                    break
            blocks.append({"type": "list", "ordered": True, "items": list_items})
            continue

        # Blockquote
        if line.startswith(">"):
            quote_lines = [line[1:].strip()]
            i += 1
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i][1:].strip())
                i += 1
            blocks.append({"type": "quote", "text": "\n".join(quote_lines)})
            continue

        # Empty line — skip
        if line.strip() == "":
            i += 1
            continue

        # Paragraph (collect until empty line or block-level element)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != "" and not (
            lines[i].startswith("#") or lines[i].startswith("```") or
            lines[i].startswith("|") or lines[i].startswith("- ") or
            lines[i].startswith("* ") or re.match(r"^\d+\.\s", lines[i]) or
            lines[i].startswith(">")
        ):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines).strip()})

    return blocks


# ─────────────────────────────────────────────────────────────────────────────
# Inline formatting: **bold**, *italic*, `code`, [link](url)
# ─────────────────────────────────────────────────────────────────────────────
INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def add_formatted_runs(paragraph, text: str):
    """Procesa **bold**, *italic*, `code`, [link](url) y añade runs al párrafo."""
    if not text:
        return
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part.startswith("[") and "](" in part:
            text_part = part[1:part.index("]")]
            run = paragraph.add_run(text_part)
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        else:
            paragraph.add_run(part)


# ─────────────────────────────────────────────────────────────────────────────
# Renderizado de bloques al .docx
# ─────────────────────────────────────────────────────────────────────────────
HEADING_SIZES = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}


def render_blocks(doc: Document, blocks: list):
    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = block["level"]
            p = doc.add_paragraph()
            p.style = doc.styles[f"Heading {min(level, 6)}"] if f"Heading {min(level, 6)}" in [s.name for s in doc.styles] else p.style
            run = p.add_run(block["text"])
            run.font.size = Pt(HEADING_SIZES.get(level, 11))
            run.font.bold = True
            run.font.name = FONT_NAME
            if level == 1:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            elif level == 2:
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        elif btype == "paragraph":
            p = doc.add_paragraph()
            add_formatted_runs(p, block["text"])
            for run in p.runs:
                run.font.name = FONT_NAME
                if not run.font.size:
                    run.font.size = Pt(11)

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            n_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=n_cols)
            table.style = "Table Grid"
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx >= n_cols:
                        break
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_formatted_runs(p, cell_text)
                    for run in p.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)
                        if row_idx == 0:  # header row
                            run.font.bold = True
                    if row_idx == 0:
                        set_cell_background(cell, "D9D9D9")
            doc.add_paragraph()  # spacing after table

        elif btype == "list":
            for item in block["items"]:
                style = "List Number" if block["ordered"] else "List Bullet"
                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:
                    p = doc.add_paragraph()
                add_formatted_runs(p, item["text"])
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(11)

        elif btype == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # background fill of paragraph (light gray)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)

        elif btype == "quote":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.italic = True
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        elif btype == "hr":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "808080")
            pBdr.append(bottom)
            pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("md_file", help="Archivo Markdown de entrada generado por doc-writer")
    parser.add_argument("--output", help="Archivo .docx de salida (default: mismo nombre con .docx)")
    parser.add_argument("--doc-type", choices=list(DOC_TYPE_CONFIG.keys()),
                        help="Tipo de documento (default: inferir del nombre)")
    parser.add_argument("--logo", help="Path al logo FEMSA Comercio (default: .claude/standards/femsa_comercio_logo.png)")
    args = parser.parse_args()

    md_path = Path(args.md_file)
    if not md_path.exists():
        sys.exit(f"ERROR: archivo no encontrado: {md_path}")

    doc_type = args.doc_type or infer_doc_type(md_path.name)
    print(f"→ Tipo de documento: {doc_type} ({DOC_TYPE_CONFIG[doc_type]['header_title']})")

    output_path = Path(args.output) if args.output else md_path.with_suffix(".docx")

    # Resolver path del logo: si no se pasó, buscar en .claude/standards/ relativo al .md
    if args.logo:
        logo_path = Path(args.logo)
    else:
        # buscar hacia arriba .claude/standards/femsa_comercio_logo.png
        candidate = md_path.parent
        logo_path = None
        for _ in range(5):
            test = candidate / ".claude" / "standards" / "femsa_comercio_logo.png"
            if test.exists():
                logo_path = test
                break
            candidate = candidate.parent
        if logo_path is None:
            # fallback al script's dir
            logo_path = Path(__file__).parent.parent / ".claude" / "standards" / "femsa_comercio_logo.png"

    # Leer markdown
    md_text = md_path.read_text(encoding="utf-8")
    blocks = parse_markdown_to_blocks(md_text)
    print(f"→ {len(blocks)} bloques parseados")

    # Construir documento
    doc = Document()

    # Configurar márgenes (similares al .dotx oficial)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header + Footer corporativos
    build_header(doc.sections[0], doc_type, logo_path)
    build_footer(doc.sections[0])

    # Renderizar contenido
    render_blocks(doc, blocks)

    # Guardar
    doc.save(str(output_path))
    print(f"✓ Generado: {output_path}")
    print(f"  Logo usado: {logo_path}")
    print(f"  Header: {DOC_TYPE_CONFIG[doc_type]['header_title']} / {DOC_TYPE_CONFIG[doc_type]['header_code']}")


if __name__ == "__main__":
    main()
